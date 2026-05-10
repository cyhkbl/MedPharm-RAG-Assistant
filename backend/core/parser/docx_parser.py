from __future__ import annotations

import uuid
from datetime import datetime
from pathlib import Path

from docx import Document

from backend.utils.text_utils import clean_text, count_chars


async def parse_docx(file_path: str) -> dict:
    """Parse a DOCX textbook and split by Heading 1/2 paragraphs."""

    path = Path(file_path)
    document = Document(str(path))
    title = _document_title(document, path)
    chapters: list[dict] = []
    current_title = title
    current_parts: list[str] = []

    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if style_name.lower().startswith(("heading 1", "heading 2")) or style_name.startswith(("标题 1", "标题 2")):
            if current_parts:
                chapters.append(_chapter(len(chapters) + 1, current_title, "\n".join(current_parts)))
            current_title = text
            current_parts = [text]
        else:
            current_parts.append(text)

    if current_parts:
        chapters.append(_chapter(len(chapters) + 1, current_title, "\n".join(current_parts)))

    total_chars = sum(chapter["char_count"] for chapter in chapters)
    return {
        "id": _textbook_id(path),
        "filename": path.name,
        "title": title,
        "format": "docx",
        "total_pages": max(len(chapters), 1),
        "total_chars": total_chars,
        "chapters": chapters,
        "upload_time": datetime.utcnow().isoformat(),
        "parse_status": "success" if chapters else "empty",
    }


def _document_title(document: Document, path: Path) -> str:
    if document.core_properties.title:
        return str(document.core_properties.title).strip()
    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)
        if text:
            return text[:120]
    return path.stem


def _chapter(index: int, title: str, content: str) -> dict:
    cleaned = clean_text(content)
    return {
        "chapter_id": f"ch_{index:03d}",
        "title": title[:120] or f"章节 {index}",
        "page_start": index,
        "page_end": index,
        "content": cleaned,
        "char_count": count_chars(cleaned),
    }


def _textbook_id(path: Path) -> str:
    return f"tb_{uuid.uuid5(uuid.NAMESPACE_URL, str(path.resolve()))}"
